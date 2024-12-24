import os
import json
# from PIL import Image
import google.generativeai as genai
from io import StringIO
import PyPDF2
import docx
from io import BytesIO
import os
import streamlit as st

import os
import streamlit as st
import tempfile
from pathlib import Path
import base64
from spire.doc import *
from spire.doc.common import *



working_dir = os.path.dirname(os.path.abspath(__file__))

config_file = f"{working_dir}/config.json"
config_data = json.load(open(config_file))

# Loading the API key

Google_Api_key = config_data["GOOGLE_API_KEY"]

# Configuring google generativeai with api kay
genai.configure(api_key=Google_Api_key)

# Function For Chatbot
def load_gemini_pro_model():
    gemini_pro_model = genai.GenerativeModel("gemini-1.5-flash")
    return gemini_pro_model

# Function For image captioning ..
def gemini_pro_vision_response(prompt,image):
    gemini_pro_vision_model = genai.GenerativeModel("gemini-1.5-flash")
    response = gemini_pro_vision_model.generate_content([prompt,image])
    result= response.text
    return result


def read_pdf_or_docx(uploaded_file):
    file_name = uploaded_file.name
    
    if file_name.endswith(".pdf"):
        reader = PyPDF2.PdfReader(uploaded_file)
        text = ""
        for page_num in range(len(reader.pages)):
            page = reader.pages[page_num]
            text += page.extract_text()
        return text
    
    elif file_name.endswith(".docx") or file_name.endswith(".doc"):
        doc = docx.Document(uploaded_file)
        text = "\n".join([para.text for para in doc.paragraphs])
        return text
    
    else:
        return "Unsupported file type"

def convert_pdf_to_html(pdf_file):
    # Create an in-memory buffer to store the HTML output
    output_buffer = BytesIO()
    # Convert the PDF to HTML and write the HTML to the buffer
    with open(pdf_file, 'rb') as pdf_file:
        extract_text_to_fp(pdf_file, output_buffer, output_type='html')
    # Retrieve the HTML content from the buffer
    html_content = output_buffer.getvalue().decode('utf-8')
    return html_content

def create_download_link(pdf_content, filename):
    """Create a download link for the PDF file"""
    b64 = base64.b64encode(pdf_content).decode()
    return f'<a href="data:application/pdf;base64,{b64}" download="{filename}">Click here to download PDF</a>'

def markdown_to_pdf(markdown_text, output_filename="resume.pdf"):
    """Convert markdown text to PDF and provide download link"""
    try:
        # write markdown to a temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix=".md") as f:
            f.write(markdown_text.encode())
            temp_md_file = f.name
        # Create a Document object
        document = Document()

        # Load a Markdown file
        document.LoadFromFile(temp_md_file)

        # Save it as a pdf file
        document.SaveToFile("ToPdf.pdf", FileFormat.PDF)
        
        # Read the generated PDF
        with open("ToPdf.pdf", "rb") as f:
            pdf_content = f.read()
        
        return pdf_content
        
    except Exception as e:
        st.error(f"Error converting markdown to PDF: {str(e)}")
        return None